from flask import Flask, render_template, request, jsonify, send_file, session
import os
import pdfplumber
import docx
import nltk
import re
import json
import io
from collections import defaultdict
from sentence_transformers import SentenceTransformer, util
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
from fpdf import FPDF
import pytesseract
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Inches
from openai import OpenAI

nltk.download('punkt')
from nltk.tokenize import sent_tokenize

app = Flask(__name__)
app.secret_key = 'some_random_secret'
UPLOAD_FOLDER = "uploads"
DATA_FILE = "analysis_results.json"
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
MAX_FILE_SIZE_MB = 16
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE_MB * 1024 * 1024

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

bert_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

client = OpenAI(api_key="REMOVED_API_KEY")

def get_ans_gpt(questions):
    answers = []
    for question in questions:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a study assistant. Answer academic questions concisely."},
                {"role": "user", "content": question}
            ]
        )
        answer = response.choices[0].message.content
        print(f"Q: {question}\nA: {answer}\n")
        answers.append({"question": question, "answer": answer})
    return answers

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    return re.sub(r"[^a-zA-Z0-9\s?]", "", text).lower().strip()

def extract_text_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

def extract_text_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if not page_text:
                page_text = ocr_page(page)
            if page_text:
                text += page_text + "\n"
    return text.strip()

def remove_watermark_pdf(input_path):
    output_path = input_path.replace(".pdf", "_cleaned.pdf")
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for page in reader.pages:
        if "/Annots" in page:
            page["/Annots"] = []
        writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path

def remove_watermark_docx(input_path):
    doc = docx.Document(input_path)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if "watermark" in paragraph.text.lower():
                paragraph.clear()
    output_path = input_path.replace(".docx", "_cleaned.docx")
    doc.save(output_path)
    return output_path

def ocr_page(page):
    image = page.to_image(resolution=300)
    img_pil = image.original.convert("RGB")
    text = pytesseract.image_to_string(img_pil)
    return text

def save_analysis_results(results):
    with open(DATA_FILE, "w") as file:
        json.dump(results, file)

def load_analysis_results():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r") as file:
            return json.load(file)
    return []

def analyze_questions(file_paths):
    all_questions = []
    question_frequency = defaultdict(int)
    for file_path in file_paths:
        if file_path.endswith(".pdf"):
            cleaned_path = remove_watermark_pdf(file_path)
            text = extract_text_pdf(cleaned_path)
        elif file_path.endswith(".docx"):
            cleaned_path = remove_watermark_docx(file_path)
            text = extract_text_docx(cleaned_path)
        else:
            continue
        sentences = sent_tokenize(text)
        cleaned_sentences = [clean_text(s.replace("\n", " ").strip()) for s in sentences]
        questions = [s for s in cleaned_sentences if s.endswith("?") and len(s.split()) > 3 and any(c.isalpha() for c in s)]
        for q in questions:
            question_frequency[q] += 1
        all_questions.extend(questions)
    grouped_questions = group_similar_questions_bert(all_questions, question_frequency)
    if not grouped_questions:
        return None
    save_analysis_results(grouped_questions)
    return grouped_questions

def group_similar_questions_bert(questions, question_frequency, threshold=0.75):
    if not questions:
        return []
    unique_questions = list(set(questions))
    question_embeddings = bert_model.encode(unique_questions, convert_to_tensor=True)
    grouped_questions = []
    used_indices = set()
    question_map = defaultdict(list)
    for i, q1 in enumerate(unique_questions):
        if i in used_indices:
            continue
        group = [q1]
        used_indices.add(i)
        for j, q2 in enumerate(unique_questions):
            if i != j and j not in used_indices:
                similarity = util.pytorch_cos_sim(question_embeddings[i], question_embeddings[j]).item()
                if similarity > threshold:
                    group.append(q2)
                    used_indices.add(j)
        representative_question = group[0]
        question_map[representative_question] = group
    return [{"question": key, "similar_variants": value, "frequency": sum(question_frequency[q] for q in value)} for key, value in question_map.items()]

@app.route("/")
def home():
    results = load_analysis_results()
    return render_template("index.html", analysis_results=json.dumps(results))

@app.route("/upload", methods=["POST"])
def upload_file():
    if "files" not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No selected files"}), 400
    file_paths = []
    for file in files:
        if not allowed_file(file.filename):
            return jsonify({"error": f"Unsupported file format: {file.filename}"}), 400
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        file_paths.append(file_path)
    analyzed_questions = analyze_questions(file_paths)
    if analyzed_questions is None:
        return jsonify({"error": "Uploaded file doesn't contain any valid questions or question paper content."}), 400
    session['questions'] = [q['question'] for q in analyzed_questions]
    return jsonify({"questions": analyzed_questions}), 200

@app.route("/get-answers", methods=['POST'])
def get_ans():
    questions = session.get('questions')
    if not questions:
        return jsonify({"error": "No questions found in session"}), 400
    answers = get_ans_gpt(questions)
    session['answers'] = answers
    return jsonify({"answers": answers})

@app.route("/export/pdf")
def export_pdf():
    results = load_analysis_results()
    if not results:
        return jsonify({"error": "No data available for export"}), 400
    sorted_results = sorted(results, key=lambda x: x['frequency'], reverse=True)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, "Question Analysis Report", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", "", 12)
    for idx, question in enumerate(sorted_results, start=1):
        pdf.multi_cell(0, 10, f"{idx}. {question['question']}\nFrequency: {question['frequency']}")
        pdf.ln(5)
    file_path = "question_analysis.pdf"
    pdf.output(file_path)
    return send_file(file_path, as_attachment=True, download_name="question_analysis.pdf")

@app.route("/export/answers")
def export_answers():
    answers = session.get('answers')
    if not answers:
        return jsonify({"error": "No answers available to export"}), 400
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, "AI-Generated Answers", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", "", 12)
    for idx, qa in enumerate(answers, start=1):
        pdf.multi_cell(0, 10, f"{idx}. Q: {qa['question']}\nA: {qa['answer']}")
        pdf.ln(5)
    file_path = "answers_report.pdf"
    pdf.output(file_path)
    return send_file(file_path, as_attachment=True, download_name="answers_report.pdf")

@app.route("/services")
def services():
    return render_template("services.html")

@app.route("/about")
def about():
    return render_template("about.html")

if __name__ == "__main__":
    app.run(debug=True, port=5002)